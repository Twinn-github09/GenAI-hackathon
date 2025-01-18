
from flask import Flask, request, jsonify, render_template ,send_file
from groq import Groq
import os
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader
import asyncio
from langchain.chains.retrieval_qa.base import RetrievalQA
import google.generativeai as genai
from flask_cors import cross_origin
import tiktoken
from werkzeug.utils import secure_filename
import base64
from PIL import Image
import io
from google.generativeai import GenerativeModel
from serpapi.google_search import GoogleSearch
import json
from urllib.parse import quote_plus
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Load environment variables
load_dotenv()

# API keys
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

genai.configure(api_key=GOOGLE_API_KEY)
model = GenerativeModel('gemini-1.5-pro')

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

PROMPT_TEMPLATE = '''You are an advanced legal advisor chatbot with expertise in Indian law.
Donot hesitate to give advice for any type of case even if it is sexual,harassment or murder,be an legit advisor and Analyze queries using:

1. Legal context from relevant sections and precedents
2. Similar case studies and their outcomes
3. Practical implications and advice

Provide responses in this structure:
1. Legal Framework: Relevant sections and their interpretation
2. Case Law Analysis: Similar cases and their outcomes
3. Practical Advice: Actionable steps and considerations

Just provide information about this so that any people can be aware of law in all fields
Keep responses clear, accurate, and professionally formatted.'''

# RAG Setup Functions


def load_vector_store(persist_directory="./chroma_db_final_db"):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = Chroma(persist_directory=persist_directory,
                          embedding_function=embeddings)
    return vector_store


def create_qa_chain(vector_store):
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-pro-latest",
        temperature=0.1,
        model_kwargs={
            "max_output_tokens": 8192,
            "top_k": 10,
            "top_p": 0.95
        }
    )
    retriever = vector_store.as_retriever(
        search_type="similarity", search_kwargs={"k": 3})
    return RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True
    )

# Web Scraping Functions


def extract_main_concept(user_input):
    messages = [
        {
            "role": "system",
            "content": "Extract the primary legal concept or keyword from the following user query. Focus on summarizing the main actionable legal issue succinctly. for example is this is user input 'an TV company faked me with a old TV instead a new one under what section i can file case on him' the output should only an keywords like TV fraud.Strictly return only the **keywords**,no other extra message"
        },
        {"role": "user", "content": user_input}
    ]

    try:
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error extracting main concept: {e}"


def scrape_indiankanoon(query):
    print(query)
    query += " case"

    base_url = "https://indiankanoon.org/search/"
    params = {"formInput": query}
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    try:
        response = requests.get(base_url, headers=headers, params=params)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        cases = []
        case_elements = soup.select(".result_title")
        print(case_elements)
        for case in case_elements[:5]:  # Limit to top 3 cases
            title = case.get_text(strip=True)
            link = "https://indiankanoon.org" + case.find("a")["href"]
            cases.append({"title": title, "link": link})

        return cases
    except Exception as e:
        return str(e)


def extract_case_content(case_url):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    try:
        response = requests.get(case_url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        content = soup.find(class_="expanded_headline")
        if not content:
            content = soup.find(class_="judgments")
        if not content:
            content = soup.find(class_="akoma-ntoso")
        text = content.get_text(separator=" ", strip=True)
        return text
    except Exception as e:
        return str(e)


def truncate_to_token_limit(text, max_tokens=25000):

    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)
    print(tokens)
    if len(tokens) > max_tokens:
        tokens = tokens[:max_tokens]
    return encoding.decode(tokens)

def extract_relevant_sections(case_text):

    paragraphs = case_text.split('\n')
    
    intro_length = min(7, len(paragraphs))
    outro_length = min(7, len(paragraphs))
    
    relevant_text = '\n'.join(paragraphs[:intro_length] + paragraphs[-outro_length:] )
    print(relevant_text)
    return relevant_text

def generate_summary_and_advice(case_text):

    relevant_text = extract_relevant_sections(case_text)
    
    messages = [
        {
            "role": "system",
            "content": """Analyze this legal case and provide a summary in Markdown format:

            # Case Summary
            ## Key Facts
            - (2-3 key facts)

            ## Legal Principles
            - (2-3 key principles)

            ## Outcome & Implications
            - (Main outcome and implications)"""
        },
        {"role": "user", "content": relevant_text}
    ]

    try:
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating summary: {e}"

# def generate_summary_and_advice(case_text):
#     messages = [
#         {
#             "role": "system",
#             "content": """Analyze this legal case and provide a summary in Markdown format:

#             # Case Summary
#             ## Key Facts
#             - (2-3 key facts)

#             ## Legal Principles
#             - (2-3 key principles)

#             ## Outcome & Implications
#             - (Main outcome and implications)"""
#         },
#         {"role": "user", "content": case_text}
#     ]

#     try:
#         response = client.chat_completions.create(
#             model="gemini-1.5-flash",
#             messages=messages
#         )
#         return response.choices[0].message.content
#     except Exception as e:
#         return f"Error generating summary: {e}"


@app.route('/chat', methods=['POST'])
@cross_origin()
def chat():
    try:
        data = request.json
        if not data or 'user_prompt' not in data:
            return jsonify({"error": "User prompt is required."}), 400

        user_prompt = data['user_prompt']

        greetings = {"hi", "hii", "hello", "hey", "hola", "namaste"}
        if user_prompt in greetings:
            return jsonify({"main_response": "Welcome to the Legal Advisor Chatbot! How can I assist you today?", "related_cases": []}), 200

        if not isinstance(user_prompt, str) or not user_prompt.strip():
            return jsonify({"error": "Invalid user prompt"}), 400

        # 1. Extract Main Concept
        main_concept = extract_main_concept(user_prompt)
        if not main_concept or "Error" in main_concept:
            return jsonify({"error": f"Failed to extract main concept: {main_concept}"}), 500

        # 2. Query RAG System
        try:
            vector_store = load_vector_store()
            qa_chain = create_qa_chain(vector_store)
            user_prompt += " If not know please return the nearest artice or section under which this case come under in law"
            rag_response = qa_chain.invoke({"query": user_prompt})
            print(rag_response["result"])
        except Exception as e:
            print(f"RAG system error: {str(e)}")
            return jsonify({"error": "Failed to process with RAG system"}), 500

        # 3. Get Groq Response
        try:
            messages = [
                {"role": "system", "content": PROMPT_TEMPLATE},
                {"role": "user", "content": f"""
                    Question: {user_prompt}
                    Relevant Legal Context: {rag_response['result']}
                    If Legal context is not given or empty consider question and get correct legal context like section in law and use that for Relevant legal context
                    Provide a comprehensive response incorporating this legal context.
                """}
            ]

            groq_response = client.chat.completions.create(
                model="llama3-8b-8192",
                messages=messages
            )
        except Exception as e:
            print(f"Groq API error: {str(e)}")
            return jsonify({"error": "Failed to get AI response"}), 500

        # 4. Get Related Cases
        try:
            cases = scrape_indiankanoon(main_concept)
            if isinstance(cases, str):
                print(f"Scraping error: {cases}")
                cases = []  # Continue with empty cases instead of failing

            detailed_cases = []
            for case in cases:
                case_content = extract_case_content(case["link"])
                if isinstance(case_content, str) and "Error" not in case_content:
                    summary_and_advice = generate_summary_and_advice(
                        case_content)
                    detailed_cases.append({
                        "title": case["title"],
                        "link": case["link"],
                        "summary_and_advice": summary_and_advice
                    })
        except Exception as e:
            print(f"Case processing error: {str(e)}")
            detailed_cases = []  # Continue with empty cases

        # 5. Compile Complete Response
        response_data = {
            "main_response": groq_response.choices[0].message.content,
            "related_cases": detailed_cases
        }

        return jsonify(response_data), 200

    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500


@app.route('/pdf_chat', methods=['POST'])
@cross_origin()
def handle_pdf():
    file = request.files['pdf']
    if file and file.filename.endswith('.pdf'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

    loader = PyPDFLoader(filepath)
    pages = loader.load()

    page_content = "\n".join([page.page_content for page in pages])

    messages = [
        {
            "role": "system",
            "content": """
            Analyze the following legal document provided by the user. Your task is to break down the document and offer a detailed analysis. Provide the response in Markdown format, covering the following sections:

            # Document Analysis

            ## Summary
            - Summarize the document briefly.
            - Highlight key points, including its main purpose and content.

            ## Legal Context
            - Identify relevant laws, regulations, or legal principles referenced or applicable.
            - Explain the context of the document in the broader legal framework.

            ## Implications
            - Detail the legal implications of the document.
            - Discuss any obligations, rights, or risks associated with it.

            ## Recommended Actions
            - Provide clear, actionable **steps** the user should take based on the document.
            - Include timelines, if applicable.
            - Suggest whether consulting a **legal professional is necessary** and why.

            ## Guidance
            - Explain the importance of the document in simple terms.
            - Offer practical considerations and general advice for the user's situation.

            Be concise, yet thorough, and give guidance to the user's perspective, ensuring accessibility and understanding.
            Give me the response with proper break lines and spaces between heading , subheading and paragraph"""
        },
        {"role": "user", "content": page_content}
    ]

    response = client.chat.completions.create(
        model="llama3-8b-8192",
        messages=messages
    )
    if filepath and os.path.exists(filepath):
        os.remove(filepath)
    response_data = {
        "main_response": response.choices[0].message.content,
        "related_cases": []
    }
    return jsonify(response_data), 200

@app.route('/process_image', methods=['POST'])
@cross_origin()
def process_image():
        
        file = request.files['image']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        # Read and process the image
        image_bytes = file.read()
        image = Image.open(io.BytesIO(image_bytes))
        
        # Convert image to base64
        buffered = io.BytesIO()
        image.save(buffered, format=image.format or 'JPEG')
        img_str = base64.b64encode(buffered.getvalue()).decode()

        # Generate caption using Gemini
        response = model.generate_content([
    {
        "inline_data": {  # Corrected key
            "data": img_str,
            "mime_type": file.content_type or "image/jpeg"  
        }
    },
    "Analyze this image and provide a detailed caption focusing on any potential legal aspects or concerns visible in the image. Return only the legal happening in the image not other messages"
])


        caption = response.text
        caption+=" Under which section in law does it come"
        print(caption)
        chat_payload = {'user_prompt': caption}
        chat_url = 'http://127.0.0.1:5000/chat'  
        chat_headers = {'Content-Type': 'application/json'}
        chat_response = requests.post(chat_url, json=chat_payload, headers=chat_headers)
        
        chat_data = chat_response.json() 
        print(chat_data.get('main_response'))
        return jsonify({
            'main_response': chat_data.get('main_response', ''),
            'related_cases': chat_data.get('related_cases', []),
        })



def convert_markdown_to_docx(markdown_text):
    doc = Document()
    
    # Parse markdown content
    lines = markdown_text.split('\n')
    current_list = []
    list_level = 0
    
    for line in lines:
        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(16 - level)  # Decrease size for deeper headers
            
        # Handle bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            paragraph = doc.add_paragraph(text)
            paragraph.style = 'List Bullet'
            
        # Handle numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            paragraph = doc.add_paragraph(text)
            paragraph.style = 'List Number'
            
        # Handle bold text
        elif '**' in line:
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
                    
        # Regular text
        else:
            doc.add_paragraph(line)
    
    return doc


def analyze_case_probability(case_details):
    try:
        # Format search query
        search_query = f"success rate in cases similar to {case_details} India legal judgments"
        encoded_query = quote_plus(search_query)
        
        # Use SerpAPI to get search results
        search = GoogleSearch({
            "q": encoded_query,
            "api_key": os.getenv("SERPAPI_API_KEY"),
            "num": 10  # Number of results to analyze
        })
        results = search.get_dict()
        
        # Extract relevant search results
        organic_results = results.get("organic_results", [])
        search_snippets = [result.get("snippet", "") for result in organic_results]
        
        # Combine all search results for analysis
        combined_text = "\n".join(search_snippets)
        print(combined_text)
        # Use Groq to analyze the probability
        messages = [
            {
                 "role": "system",
                "content": """Analyze the provided case summaries to determine:
                - Probability of winning based on precedents
                - Key factors and challenges
                - Suggestions for improvement
                
                Return the response in this format:
                {
                    "success_probability": "X%",
                    "confidence_level": "High/Medium/Low",
                    "key_factors": ["factor1", "factor2", ...],
                    "challenges": ["challenge1", "challenge2", ...],
                    "recommendation": "brief recommendation"
                }
                """
            },
            {
                "role": "user",
                "content": f"Case Details: {case_details}\n\nSimilar Cases:\n{combined_text}"
            }
        ]
        
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages
        )
        print("---------------")
        print(response)
        # Parse the response
        analysis = response.choices[0].message.content
        print(analysis)
        return analysis
        
    except Exception as e:
        print(f"Error in case analysis: {str(e)}")
        return {
            "success_probability": "Unknown",
            "confidence_level": "Low",
            "key_factors": ["Insufficient data for analysis"],
            "challenges": ["Unable to analyze similar cases"],
            "recommendation": "Consult with a legal professional for detailed analysis"
        }
        
@app.route('/get-help' ,methods=['POST'])
@cross_origin()
def handle_help():
    
    data = request.json
    if not data or 'user_prompt' not in data:
        return jsonify({"error": "User prompt is required."}), 400

    user_prompt = data['user_prompt']
    print(user_prompt)
    case_analysis = analyze_case_probability(user_prompt)
    messages = [
                {"role": "user", "content": f"""
                 Case details : {user_prompt}
                 The user intends to create a new petition or file a case. Your task is to guide them by providing a clear and proper format for drafting the document. Ensure the format aligns with legal standards and is tailored to the user's needs. 

                Provide the response in Markdown format, including the following sections:

                # Petition/Case Filing Template

                ## General Instructions
                - Briefly explain the purpose of the document.
                - Highlight any key details the user must provide (e.g., names, dates, case details).

                ## Petition Format
                Provide a standard structure, such as:
                1. **Header**:
                - Title of the petition or case filing.
                - Court details (name, jurisdiction, etc.).
                - Names of the petitioner(s) and respondent(s).
                2. **Introduction**:
                - Brief background of the case.
                - Purpose of the petition.
                3. **Statement of Facts**:
                - Chronological summary of events leading to the petition.
                4. **Legal Grounds**:
                - Relevant laws and principles supporting the petition.
                5. **Prayer for Relief**:
                - Specific outcomes or remedies the petitioner is seeking.
                6. **Declaration**:
                - Statement of truth signed by the petitioner.
                7. **Annexures** (if applicable):
                - List supporting documents attached to the petition.

                ## Customizable Sections
                - Highlight areas the user should fill in with their details.
                - Provide examples or sample text for commonly used phrases.

                ## Practical Guidance
                - Suggest steps to finalize and submit the document.
                - Advise if legal assistance is recommended.

                Ensure the template is adaptable to various legal scenarios (e.g., civil cases, criminal cases, appeals).
                The output should be with a include name address and other details of the user in the template.
                Just give me the template donot add extra message, like "here is your answer etc..."
                At last
                The template should include standard legal formatting and also highlight:
                based on {case_analysis}
                include success rate
                1. The strength of similar precedents 
                2. Key factors that support the case
                3. Try incluing all kind of postive and negative factors where user face in the case
                
                Provide the response in Markdown format with all standard petition sections.
                Include specific sections that address the identified success factors and challenges.
                """}
            ]

    groq_response = client.chat.completions.create(
                model="llama3-8b-8192",
                messages=messages
            )
    markdown_content = groq_response.choices[0].message.content
    
    doc = convert_markdown_to_docx(markdown_content)
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    import base64
    docx_base64 = base64.b64encode(docx_buffer.getvalue()).decode()
    
    response_data = {
        "main_response": markdown_content,
        "docx_content": docx_base64,
        "related_cases": []
    }

    return jsonify(response_data), 200
    
@app.route('/download-docx', methods=['POST'])
def download_docx():
    data = request.json
    if not data or 'markdown_content' not in data:
        return jsonify({"error": "Markdown content is required."}), 400
        
    doc = convert_markdown_to_docx(data['markdown_content'])
    
    # Save to bytes buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return send_file(
        docx_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='petition_template.docx'
    )


@app.route('/explain-legal-term', methods=['POST'])
@cross_origin()
def explain_legal_term():
    # Extract the user input
    data = request.json
    if not data or 'user_prompt' not in data:
        return jsonify({"error": "User prompt is required."}), 400

    user_prompt = data['user_prompt']
    print(f"User prompt received: {user_prompt}")

    # Prepare messages for the AI model
    messages = [
        {"role": "user", "content": f"""
        The user wants to understand the legal term: **{user_prompt}**.
        Your task is to provide a clear, concise, and detailed explanation of the term. Ensure the explanation is understandable to someone without a legal background and includes the following sections:

        # Explanation of {user_prompt}

        ## Definition
        - Provide a formal definition of the term.
        - Highlight its importance in the legal context.

        ## Legal Context and Usage
        - Explain where this term is typically used (e.g., contracts, court cases, statutes).
        - Include relevant examples to illustrate its application.

        ## Related Legal Terms
        - Mention any terms closely related or commonly associated with the provided term.
        - Briefly explain how they are connected.

        ## Practical Implications
        - Describe why it’s important for someone to understand this term.
        - Mention any common misconceptions or pitfalls associated with it.

        ## Additional Notes (if applicable)
        - Add any historical background, legal doctrines, or notable cases involving this term.
        - Include references to relevant laws or legal texts if necessary.

        Ensure the explanation is detailed yet accessible. Avoid using unnecessary legal jargon unless it's essential, and explain it where used. Provide the output in Markdown format.
        """}
    ]

    # Query the AI model
    try:
        groq_response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages
        )
        print(f"AI response: {groq_response}")

        # Prepare response data
        response_data = {
            "main_response": groq_response.choices[0].message.content,
        }
        return jsonify(response_data), 200

    except Exception as e:
        print(f"Error occurred: {e}")
        return jsonify({"error": "An error occurred while processing the request."}), 500


@app.route('/gender-neutral-language-suggestions', methods=['POST'])
@cross_origin()
def gender_neutral_language_suggestions():
    # Extract the user input
    data = request.json
    if not data or 'user_text' not in data:
        return jsonify({"error": "Text input is required."}), 400

    user_text = data['user_text']
    print(f"User text received: {user_text}")

    # Prepare messages for the AI model
    messages = [
        {"role": "user", "content": f"""
        The user wants to ensure gender-neutral language in their text: **{user_text}**.
        Your task is to detect gendered language and suggest neutral alternatives. 

        # Gender-Neutral Language Suggestions

        ## Gendered Terms Detected
        - Identify any gendered terms or phrases (e.g., "he", "she", "chairman") in the text.
        - Suggest gender-neutral alternatives for each term (e.g., "they", "chairperson").

        ## Suggestions Format
        - Return the list of detected gendered terms and their neutral replacements.
        - Provide the new text with the suggested replacements integrated.

        ## Example:
        - Original: "The chairman will make his decision soon."
        - Suggested: "The chairperson will make their decision soon."

        Ensure the suggestions are clear, and keep the integrity of the original meaning while avoiding gender-specific language.
        """}
    ]

    # Query the AI model
    try:
        groq_response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages
        )
        print(f"AI response: {groq_response}")

        # Prepare response data
        response_data = {
            "main_response": groq_response.choices[0].message.content,
        }
        return jsonify(response_data), 200

    except Exception as e:
        print(f"Error occurred: {e}")
        return jsonify({"error": "An error occurred while processing the request."}), 500
    
    
@app.route('/nearby-lawyers', methods=['POST'])
@cross_origin()
def get_nearby_lawyers():
    try:
        
        data = request.json
        print(data)
        latitude = data.get('latitude')
        longitude = data.get('longitude')
        
        url = 'https://maps.googleapis.com/maps/api/place/nearbysearch/json'
        
        params = {
            'location': f'{latitude},{longitude}',
            'radius': '5000',  # 5km radius
            'type': 'lawyer',
            'keyword': 'lawyer law firm attorney',
            'key': 'AIzaSyDrGnW9kUQQiNWvVk_HrqD5lVvc3I90334'
        }
        
        response = requests.get(url, params=params)
        results = response.json()
        
        lawyers = []
        for place in results.get('results', []):
            lawyer = {
                'name': place.get('name'),
                'address': place.get('vicinity'),
                'rating': place.get('rating'),
                'total_ratings': place.get('user_ratings_total'),
                'place_id': place.get('place_id'),
                'location': place.get('geometry', {}).get('location', {}),
                'open_now': place.get('opening_hours', {}).get('open_now'),
            }
            lawyers.append(lawyer)
        
        lawyers.sort(key=lambda x: (x['rating'] if x['rating'] is not None else 0), reverse=True)
       


       
        return jsonify({
            'status': 'success',
            'data': lawyers
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500
        
@app.route('/')
def home():
    return render_template('bot.html')


if __name__ == '__main__':
    app.run(debug=True)
